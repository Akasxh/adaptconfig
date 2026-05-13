"""Document upload and parsing routes."""

import asyncio
import logging
from pathlib import Path, PurePosixPath

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Query, UploadFile
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from finspark.api.dependencies import (
    get_audit_service,
    get_document_parser,
    get_tenant_context,
    require_role,
)
from finspark.core import events
from finspark.core.audit import AuditService
from finspark.core.config import settings

logger = logging.getLogger(__name__)
from finspark.core.database import get_db
from finspark.models.document import Document
from finspark.schemas.common import APIResponse, DocType, TenantContext
from finspark.schemas.documents import (
    DocumentDetailResponse,
    DocumentUploadResponse,
    ParsedDocumentResult,
)
from finspark.services.parsing.document_parser import DocumentParser
from finspark.services.webhook_delivery import deliver_event

router = APIRouter(prefix="/documents", tags=["Documents"])

ALLOWED_EXTENSIONS = {".docx", ".pdf", ".yaml", ".yml", ".json"}


@router.post("/upload", response_model=APIResponse[DocumentUploadResponse])
async def upload_document(
    file: UploadFile,
    background_tasks: BackgroundTasks,
    doc_type: str = "brd",
    db: AsyncSession = Depends(get_db),
    tenant: TenantContext = require_role("admin", "editor"),
    parser: DocumentParser = Depends(get_document_parser),
    audit: AuditService = Depends(get_audit_service),
) -> APIResponse[DocumentUploadResponse]:
    """Upload and parse a document (BRD, SOW, API spec)."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    # Sanitize filename to prevent path traversal
    safe_name = PurePosixPath(file.filename).name
    if not safe_name:
        raise HTTPException(status_code=400, detail="Invalid filename")

    suffix = Path(safe_name).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {suffix}. Allowed: {ALLOWED_EXTENSIONS}",
        )

    # Validate doc_type against the DocType enum
    try:
        DocType(doc_type)
    except ValueError:
        valid = [e.value for e in DocType]
        raise HTTPException(
            status_code=400,
            detail=f"Invalid doc_type '{doc_type}'. Allowed: {valid}",
        )

    # Validate file size before writing to disk
    max_bytes = settings.max_upload_size_mb * 1024 * 1024
    file_bytes = await file.read()
    if len(file_bytes) > max_bytes:
        raise HTTPException(
            status_code=413,
            detail=f"File exceeds maximum upload size of {settings.max_upload_size_mb} MB",
        )

    # Save file
    upload_dir = settings.upload_dir / tenant.tenant_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    file_path = upload_dir / safe_name

    def _write_file(path: Path, data: bytes) -> None:
        with open(path, "wb") as f:
            f.write(data)

    await asyncio.to_thread(_write_file, file_path, file_bytes)

    # Create document record
    doc = Document(
        tenant_id=tenant.tenant_id,
        filename=safe_name,
        file_type=suffix.lstrip("."),
        file_size=file_path.stat().st_size,
        doc_type=doc_type,
        status="parsing",
    )
    db.add(doc)
    await db.flush()

    # Extract raw text — different strategy per file type
    file_ext = suffix.lstrip(".")
    if file_ext in ("yaml", "yml", "json"):
        raw_text = file_bytes.decode("utf-8", errors="replace")
    elif file_ext == "pdf":
        from pypdf import PdfReader
        import io
        reader = PdfReader(io.BytesIO(file_bytes))
        raw_text = "\n".join(p.extract_text() or "" for p in reader.pages)
    elif file_ext == "docx":
        from docx import Document as DocxDocument
        import io
        docx_doc = DocxDocument(io.BytesIO(file_bytes))
        raw_text = "\n".join(p.text for p in docx_doc.paragraphs if p.text.strip())
    else:
        raw_text = ""

    use_llm = settings.ai_enabled and (bool(settings.openai_api_key) or bool(settings.gemini_api_key))

    # Quick regex parse (fast) — store as interim result
    regex_result = await asyncio.to_thread(parser.parse, file_path, doc_type=doc_type)
    doc.parsed_result = regex_result.model_dump_json()
    doc.raw_text = regex_result.summary[:5000]

    if use_llm:
        # LLM parsing is slow — run in background, return immediately with "parsing"
        doc_id = doc.id
        tenant_id = tenant.tenant_id
        tenant_name = tenant.tenant_name
        _bg_raw = raw_text or regex_result.summary
        _bg_safe_name = safe_name
        _bg_doc_type = doc_type
        _bg_regex_result = regex_result
        _bg_parser = parser

        async def _parse_in_background() -> None:
            from finspark.core.database import async_session_factory

            try:
                from finspark.services.llm.client import get_llm_client
                llm_client = get_llm_client()
                result = await _bg_parser.parse_with_llm(_bg_raw, _bg_safe_name, llm_client)
                logger.info("document_parsed_via_llm filename=%s", _bg_safe_name)
            except Exception:
                logger.warning("LLM parsing failed for %s, using regex", _bg_safe_name, exc_info=True)
                result = _bg_regex_result

            try:
                async with async_session_factory() as bg_db:
                    stmt = select(Document).where(Document.id == doc_id)
                    row = (await bg_db.execute(stmt)).scalar_one_or_none()
                    if row:
                        row.parsed_result = result.model_dump_json()
                        row.raw_text = result.summary[:5000]
                        row.status = "parsed"
                        await bg_db.commit()
                        logger.info("background_parse_complete doc_id=%s", doc_id)
            except Exception:
                logger.error("background_parse_db_failed doc_id=%s", doc_id, exc_info=True)

            try:
                wd = {"tenant_id": tenant_id, "document_id": doc_id, "filename": _bg_safe_name, "doc_type": _bg_doc_type}
                await events.emit(events.DOCUMENT_PARSED, wd)
                await deliver_event(tenant_id, events.DOCUMENT_PARSED, wd)
            except Exception:
                logger.warning("background_webhook_failed doc_id=%s", doc_id, exc_info=True)

        asyncio.create_task(_parse_in_background())
    else:
        # No LLM — regex is already done, mark as parsed
        doc.status = "parsed"

    # Commit so the document is visible immediately for polling
    await db.commit()

    webhook_data = {
        "tenant_id": tenant.tenant_id,
        "document_id": doc.id,
        "filename": doc.filename,
        "doc_type": doc_type,
    }
    await events.emit(events.DOCUMENT_UPLOADED, webhook_data)
    background_tasks.add_task(deliver_event, tenant.tenant_id, events.DOCUMENT_UPLOADED, webhook_data)

    await audit.log(
        tenant_id=tenant.tenant_id,
        actor=tenant.tenant_name,
        action="upload_document",
        resource_type="document",
        resource_id=doc.id,
        details={"filename": safe_name, "doc_type": doc_type, "status": doc.status},
    )

    return APIResponse(
        success=True,
        data=DocumentUploadResponse(
            id=doc.id,
            filename=doc.filename,
            file_type=suffix.lstrip("."),
            doc_type=doc_type,
            status=doc.status,
            created_at=doc.created_at,
        ),
        message="Document uploaded, parsing in progress"
        if doc.status == "parsing"
        else f"Document {doc.status}",
    )


@router.get("/{document_id}", response_model=APIResponse[DocumentDetailResponse])
async def get_document(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    tenant: TenantContext = Depends(get_tenant_context),
) -> APIResponse[DocumentDetailResponse]:
    """Get document details and parsing results."""
    stmt = select(Document).where(
        Document.id == document_id,
        Document.tenant_id == tenant.tenant_id,
    )
    result = await db.execute(stmt)
    doc = result.scalar_one_or_none()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    parsed = None
    if doc.parsed_result:
        parsed = ParsedDocumentResult.model_validate_json(doc.parsed_result)

    return APIResponse(
        data=DocumentDetailResponse(
            id=doc.id,
            filename=doc.filename,
            file_type=doc.file_type,
            doc_type=doc.doc_type,
            status=doc.status,
            parsed_result=parsed,
            created_at=doc.created_at,
            updated_at=doc.updated_at,
        ),
    )


@router.post("/{document_id}/reanalyze", response_model=APIResponse[dict])
async def reanalyze_document(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    tenant: TenantContext = require_role("admin", "editor"),
) -> APIResponse[dict]:
    """Re-run the LLM parser on an existing document.

    Used after the parser is extended (e.g., to extract chain metadata
    like depends_on/extract/inject). Existing documents won't have the
    new fields until re-parsed; this is the on-demand path so users only
    burn LLM budget on docs they care about.
    """
    stmt = select(Document).where(
        Document.id == document_id,
        Document.tenant_id == tenant.tenant_id,
    )
    result = await db.execute(stmt)
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if not settings.ai_enabled or not (settings.openai_api_key or settings.gemini_api_key):
        raise HTTPException(
            status_code=400,
            detail="LLM is not configured — re-analysis requires an LLM provider.",
        )

    # Pull the raw file from disk so we re-parse from the source, not the prior summary.
    file_path = settings.upload_dir / tenant.tenant_id / doc.filename
    if not file_path.exists():
        raise HTTPException(
            status_code=410,
            detail="Original file no longer available — re-upload the document.",
        )

    try:
        raw_bytes = file_path.read_bytes()
        ext = file_path.suffix.lower().lstrip(".")
        if ext in ("yaml", "yml", "json"):
            raw_text = raw_bytes.decode("utf-8", errors="replace")
        elif ext == "pdf":
            from pypdf import PdfReader
            import io
            raw_text = "\n".join(p.extract_text() or "" for p in PdfReader(io.BytesIO(raw_bytes)).pages)
        elif ext == "docx":
            from docx import Document as DocxDocument
            import io
            raw_text = "\n".join(p.text for p in DocxDocument(io.BytesIO(raw_bytes)).paragraphs if p.text.strip())
        else:
            raw_text = raw_bytes.decode("utf-8", errors="replace")
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"Could not read source file: {exc}") from exc

    doc.status = "parsing"
    await db.flush()

    doc_id = doc.id
    tenant_id = tenant.tenant_id
    filename = doc.filename
    doc_type = doc.doc_type

    async def _reparse_in_background() -> None:
        from finspark.core.database import async_session_factory
        from finspark.services.llm.client import get_llm_client

        try:
            llm_client = get_llm_client()
            result = await parser.parse_with_llm(raw_text, filename, llm_client)
            logger.info("document_reanalyzed_via_llm filename=%s", filename)
        except Exception:  # noqa: BLE001
            logger.warning("LLM reanalyze failed for %s, keeping prior parse", filename, exc_info=True)
            return

        async with async_session_factory() as bg_db:
            stmt2 = select(Document).where(Document.id == doc_id)
            row = (await bg_db.execute(stmt2)).scalar_one_or_none()
            if row:
                row.parsed_result = result.model_dump_json()
                row.raw_text = result.summary[:5000]
                row.status = "parsed"
                await bg_db.commit()
                logger.info("reanalyze_complete doc_id=%s", doc_id)

        try:
            wd = {"tenant_id": tenant_id, "document_id": doc_id, "filename": filename, "doc_type": doc_type}
            await events.emit(events.DOCUMENT_PARSED, wd)
            await deliver_event(tenant_id, events.DOCUMENT_PARSED, wd)
        except Exception:  # noqa: BLE001
            logger.warning("reanalyze_webhook_failed doc_id=%s", doc_id, exc_info=True)

    asyncio.create_task(_reparse_in_background())

    return APIResponse(
        data={"id": doc.id, "status": "parsing"},
        message="Re-analysis started — refresh in a few seconds.",
    )


@router.delete("/{document_id}", response_model=APIResponse[dict])
async def delete_document(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    tenant: TenantContext = require_role("admin", "editor"),
    audit: AuditService = Depends(get_audit_service),
) -> APIResponse[dict]:
    """Delete a document and its uploaded file."""
    stmt = select(Document).where(
        Document.id == document_id,
        Document.tenant_id == tenant.tenant_id,
    )
    result = await db.execute(stmt)
    doc = result.scalar_one_or_none()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    filename = doc.filename
    file_path = settings.upload_dir / tenant.tenant_id / filename
    if file_path.exists():
        file_path.unlink()

    await db.delete(doc)
    await db.flush()

    await audit.log(
        tenant_id=tenant.tenant_id,
        actor=tenant.tenant_name,
        action="delete_document",
        resource_type="document",
        resource_id=document_id,
        details={"filename": filename},
    )

    return APIResponse(
        data={"id": document_id, "deleted": True},
        message=f"Document '{filename}' deleted",
    )


@router.get("/", response_model=APIResponse[list[DocumentUploadResponse]])
async def list_documents(
    db: AsyncSession = Depends(get_db),
    tenant: TenantContext = Depends(get_tenant_context),
    page: int | None = Query(None, ge=1, description="Page number (1-based). Omit for all results."),
    page_size: int | None = Query(None, ge=1, le=200, description="Items per page. Omit for all results."),
) -> APIResponse[list[DocumentUploadResponse]]:
    """List all documents for the current tenant."""
    stmt = (
        select(Document)
        .where(Document.tenant_id == tenant.tenant_id)
        .order_by(Document.created_at.desc())
    )
    if page is not None and page_size is not None:
        stmt = stmt.offset((page - 1) * page_size).limit(page_size)
    result = await db.execute(stmt)
    docs = result.scalars().all()

    return APIResponse(
        data=[
            DocumentUploadResponse(
                id=d.id,
                filename=d.filename,
                file_type=d.file_type,
                doc_type=d.doc_type,
                status=d.status,
                created_at=d.created_at,
            )
            for d in docs
        ],
    )
